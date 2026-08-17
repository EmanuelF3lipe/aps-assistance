#!/usr/bin/env python3
"""Gera a documentação técnica do APS Assistance em formato .docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Estilos ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(1)

# ── Helpers ──────────────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def code(text):
    p = doc.add_paragraph(text, style='CodeBlock')
    return p

def bold_para(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)

# ════════════════════════════════════════════════════════════
# TÍTULO
# ════════════════════════════════════════════════════════════
title = doc.add_heading('APS Assistance — Documentação Técnica', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versão 2.0.0')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x90, 0xa8)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 1. STACK
# ════════════════════════════════════════════════════════════
doc.add_heading('1. Stack Tecnológica', level=1)

doc.add_paragraph(
    'O APS Assistance é um aplicativo web full-stack com arquitetura monolítica. '
    'O frontend é um SPA (Single Page Application) em React, e o backend é um servidor Express '
    'que serve tanto a API REST quanto os arquivos estáticos compilados do frontend.'
)

add_table(
    ['Camada', 'Tecnologia', 'Versão', 'Referência'],
    [
        ['Frontend Framework', 'React', '^18.2.0', 'package.json:16'],
        ['Build Tool', 'Vite', '^5.0.0', 'package.json:8'],
        ['Backend Framework', 'Express', '^4.18.2', 'package.json:12'],
        ['Real-time', 'Socket.IO (server + client)', '^4.8.3', 'package.json:13,20'],
        ['Bot Telegram', 'node-telegram-bot-api', '^1.2.0', 'package.json:14'],
        ['Ícones', 'react-icons', '^5.7.0', 'package.json:18'],
        ['Dev Runner', 'concurrently', '^8.2.2', 'package.json:10'],
    ],
    col_widths=[4, 5, 3, 4]
)

# ════════════════════════════════════════════════════════════
# 2. ESTRUTURA DE DIRETÓRIOS
# ════════════════════════════════════════════════════════════
doc.add_heading('2. Estrutura de Diretórios', level=1)

tree = """pas-interface/
├── src/                         # Frontend React
│   ├── main.jsx                 # Ponto de entrada (monta <App /> no #root)
│   ├── App.jsx                  # Componente raiz — orquestra todo o estado e navegação
│   ├── services/
│   │   └── api.js               # Cliente HTTP — todas as chamadas à API REST
│   ├── styles/
│   │   └── global.css           # Estilos globais + temas (dark/light)
│   ├── data/
│   │   └── fiscalTables.js      # Tabelas fiscais (CFOP, CST PIS/COFINS/IPI/ICMS)
│   └── components/
│       ├── SplashScreen.jsx     # Seletor de módulo (APS / Diário / Ferramentas)
│       ├── Header.jsx           # Barra superior (busca, navegação, tema)
│       ├── Sidebar.jsx          # Painel lateral (pastas, favoritos, lixeira)
│       ├── FilePanel.jsx        # Lista central de erros (ordenção, lote, CSV)
│       ├── Dashboard.jsx        # Visão geral com estatísticas
│       ├── DiarioPanel.jsx      # Módulo diário de ocorrências
│       ├── ToolboxPanel.jsx     # Módulo ferramentas (SEFAZ, CNPJ, NFe)
│       ├── ErrorPopup.jsx       # Popup de detalhe/edição do erro
│       ├── PublicForm.jsx       # Formulário de criação de erro (modal React)
│       ├── TrashPanel.jsx       # Lixeira (restaurar, excluir permanentemente)
│       ├── TagsPanel.jsx        # Navegador de tags
│       ├── RelatoriosPanel.jsx  # Gerenciamento de relatórios
│       ├── AdvancedSearchPanel.jsx # Busca avançada multi-critério
│       ├── PasswordPanel.jsx    # Painel de senhas
│       ├── Toast.jsx            # Notificações temporárias
│       └── Modals/              # Modais (pastas, arquivos, mover)
├── server/                      # Backend
│   ├── index.js                 # Servidor Express + Socket.IO (1398 linhas)
│   └── telegram-bot.js          # Bot Telegram (390 linhas)
├── public/                      # Assets estáticos (favicon, logo, cadastrar.html)
├── dist/                        # Build de produção (gerado por vite build)
├── notion/                      # Dados (herda do diretório pai — ../notion/)
├── Iniciar.bat                  # Launcher Windows
├── vite.config.js               # Configuração do Vite
└── package.json"""

for line in tree.split('\n'):
    code(line)

# ════════════════════════════════════════════════════════════
# 3. ARQUITETURA DO FRONTEND
# ════════════════════════════════════════════════════════════
doc.add_heading('3. Arquitetura do Frontend', level=1)

doc.add_heading('3.1 Ponto de entrada', level=2)
doc.add_paragraph(
    'O app inicia em src/main.jsx, que monta o componente <App /> dentro de <React.StrictMode>. '
    'O HTML de suporte é o index.html, que contém apenas o <div id="root">.'
)
code('src/main.jsx:15-18')

doc.add_heading('3.2 Componente raiz — App.jsx', level=2)
doc.add_paragraph(
    'App.jsx é o orquestrador central do application. Toda a aplicação é gerenciada por um '
    'componente funcional com múltiplos useState hooks — não há uso de Redux, Zustand ou Context API.'
)

doc.add_heading('Estado global (App.jsx:29-60)', level=3)
doc.add_paragraph(
    'O estado é dividido em categorias lógicas:'
)
add_table(
    ['Categoria', 'Variáveis de estado', 'Linha'],
    [
        ['Dados', 'folders, files, favorites, allTags', 'App.jsx:31-34'],
        ['Navegação', 'currentFolder, currentModule, mainView', 'App.jsx:36-38'],
        ['UI toggles', 'showTrash, showTags, showSplash, showNewForm, showAdvancedSearch', 'App.jsx:40-45'],
        ['Tema', 'theme (dark/light, persistido em localStorage)', 'App.jsx:47'],
        ['Seleção', 'selectedFiles (operações em lote), errorPopup (popup de detalhe)', 'App.jsx:49-50'],
    ],
    col_widths=[3, 7, 3]
)

doc.add_heading('Módulos (App.jsx:406-472)', level=3)
doc.add_paragraph(
    'O app possui três módulos principais, selecionados pelo SplashScreen:'
)
add_table(
    ['Módulo', 'currentModule', 'Componente principal', 'Descrição'],
    [
        ['APS', 'null', 'Dashboard / FilePanel / ErrorPopup', 'Catálogo de erros — sub-views: dashboard, erros, relatórios, tags, lixeira'],
        ['Diário', '"diario"', 'DiarioPanel', 'Registro de ocorrências do dia a dia'],
        ['Ferramentas', '"ferramentas"', 'ToolboxPanel', 'Utilitários fiscais (SEFAZ, CNPJ, NFe, códigos)'],
    ],
    col_widths=[3, 3, 5, 5]
)

doc.add_heading('3.3 Componentes e seus papéis', level=2)

components = [
    ['SplashScreen.jsx', 'Tela de seleção de módulo com 3 cards animados (fade-in). Renderiza fora do .app (App.jsx:375-377).'],
    ['Header.jsx', 'Barra superior: campo de busca (Ctrl+K), toggles de visão (dashboard/erros/relatórios), botão de tema, botão de atalhos, toggle da sidebar.'],
    ['Sidebar.jsx', 'Painel lateral esquerdo: lista de favoritos, lista de pastas com contagem de erros, botão de lixeira.'],
    ['FilePanel.jsx', 'Lista central de erros: ordenação, seleção em lote, exportação CSV, exclusão/movimentação em lote.'],
    ['Dashboard.jsx', 'Visão geral com cards de estatísticas: total de erros, erros por pasta, erros recentes, tags mais usadas.'],
    ['DiarioPanel.jsx', 'CRUD completo do diário: categorias (Sistema, Rede, Impressora, etc.), prioridades, turnos, autores, busca e filtros.'],
    ['ToolboxPanel.jsx', 'Ferramentas: abas (Códigos, SEFAZ, CNPJ, NFe, Referências), cada uma com funcionalidade específica.'],
    ['ErrorPopup.jsx', 'Popup modal de detalhe: exibe conteúdo Markdown, permite editar, mover, renomear, excluir, gerenciar anexos.'],
    ['PublicForm.jsx', 'Formulário React para criar novos erros (modal). Não deve ser confundido com /cadastrar (HTML standalone).'],
    ['TrashPanel.jsx', 'Visualização da lixeira: restaurar, excluir permanentemente, esvaziar lixeira.'],
    ['TagsPanel.jsx', 'Navegador de tags com contagem de erros por tag.'],
    ['AdvancedSearchPanel.jsx', 'Busca avançada: texto, pasta, tags, período.'],
    ['RelatoriosPanel.jsx', 'CRUD de relatórios pré-definidos.'],
    ['Modals/', 'NewFolderModal, NewFileModal, RenameFolderModal, MoveFileModal — diálogos para operações CRUD.'],
]

add_table(
    ['Componente', 'Função'],
    components,
    col_widths=[4.5, 12]
)

doc.add_heading('3.4 Atalhos de teclado (App.jsx:346-361)', level=2)
add_table(
    ['Atalho', 'Ação', 'Linha'],
    [
        ['Ctrl+K', 'Focar campo de busca', 'App.jsx:348'],
        ['Ctrl+N', 'Abrir formulário de novo erro', 'App.jsx:350'],
        ['Esc', 'Fechar popup, painéis, formulários', 'App.jsx:352-358'],
        ['Ctrl+/', 'Mostrar modal de atalhos', 'App.jsx:360'],
    ],
    col_widths=[3, 7, 3]
)

# ════════════════════════════════════════════════════════════
# 4. COMUNICAÇÃO FRONT ↔ BACK
# ════════════════════════════════════════════════════════════
doc.add_heading('4. Comunicação Frontend ↔ Backend', level=1)

doc.add_heading('4.1 Cliente HTTP — api.js', level=2)
doc.add_paragraph(
    'Todas as chamadas à API REST são centralizadas em src/services/api.js (384 linhas). '
    'O arquivo exporta um único objeto api com métodos agrupados por recurso.'
)
code('src/services/api.js:9 — const BASE = "/api"')
doc.add_paragraph(
    'Cada método faz fetch() com Content-Type: application/json. '
    'Não há autenticação, interceptors ou tratamento centralizado de erros.'
)

doc.add_heading('Exemplo de chamada (api.js:14-17)', level=3)
code('''async getFolders() {
    const res = await fetch(`${BASE}/folders`);
    return res.json();
}''')

doc.add_heading('Mapeamento frontend → backend (api.js)', level=3)

api_map = [
    ['api.getFolders()', 'GET /api/folders', 'api.js:14-17'],
    ['api.createFolder(name)', 'POST /api/folders', 'api.js:19-26'],
    ['api.renameFolder(old, new)', 'PUT /api/folders/:oldName', 'api.js:28-35'],
    ['api.deleteFolder(name)', 'DELETE /api/folders/:name', 'api.js:37-42'],
    ['api.getFiles(folder)', 'GET /api/files/:folder', 'api.js:47-50'],
    ['api.getFile(folder, file)', 'GET /api/file/:folder/:filename', 'api.js:52-55'],
    ['api.createFile(f, file, content)', 'POST /api/file/:folder', 'api.js:57-64'],
    ['api.updateFile(f, file, content)', 'PUT /api/file/:folder/:filename', 'api.js:66-73'],
    ['api.deleteFile(folder, file)', 'DELETE /api/file/:folder/:filename', 'api.js:75-80'],
    ['api.renameFile(f, file, new)', 'PUT /api/file/:folder/:filename/rename', 'api.js:82-89'],
    ['api.updateTags(f, file, tags)', 'PUT /api/file/:folder/:filename/tags', 'api.js:91-98'],
    ['api.moveFile(f, file, target)', 'PUT /api/file/:folder/:filename/move', 'api.js:100-107'],
    ['api.uploadAttachment(...)', 'POST /api/attachments/upload', 'api.js:111-118'],
    ['api.getAttachments(f, file)', 'GET /api/attachments/:folder/:filename', 'api.js:120-123'],
    ['api.deleteAttachment(f, file)', 'DELETE /api/attachments/:folder/:fileName', 'api.js:125-130'],
    ['api.search(query)', 'GET /api/search?q=...', 'api.js:135-138'],
    ['api.advancedSearch(filters)', 'GET /api/search/advanced?...', 'api.js:140-149'],
    ['api.getStats()', 'GET /api/stats', 'api.js:153-156'],
    ['api.getTrash()', 'GET /api/trash', 'api.js:160-163'],
    ['api.restoreFromTrash(file)', 'PUT /api/trash/restore/:filename', 'api.js:165-170'],
    ['api.deleteFromTrash(file)', 'DELETE /api/trash/:filename', 'api.js:172-177'],
    ['api.emptyTrash()', 'DELETE /api/trash', 'api.js:179-184'],
    ['api.getTags()', 'GET /api/tags', 'api.js:188-191'],
    ['api.getFavorites()', 'GET /api/favorites', 'api.js:195-198'],
    ['api.addFavorite(file, folder)', 'POST /api/favorites', 'api.js:200-207'],
    ['api.removeFavorite(file, folder)', 'DELETE /api/favorites/:folder/:filename', 'api.js:209-214'],
    ['api.getReports()', 'GET /api/reports', 'api.js:218-221'],
    ['api.createReport(...)', 'POST /api/reports', 'api.js:223-230'],
    ['api.updateReport(...)', 'PUT /api/reports/:id', 'api.js:232-239'],
    ['api.deleteReport(id)', 'DELETE /api/reports/:id', 'api.js:241-246'],
    ['api.getFolderColors()', 'GET /api/folder-colors', 'api.js:250-253'],
    ['api.setFolderColor(f, color)', 'PUT /api/folder-colors/:folder', 'api.js:255-262'],
    ['api.removeFolderColor(f)', 'DELETE /api/folder-colors/:folder', 'api.js:264-269'],
    ['api.getBotConfig()', 'GET /api/bot-config', 'api.js:273-276'],
    ['api.setBotConfig(token)', 'PUT /api/bot-config', 'api.js:278-285'],
    ['api.stopBot()', 'POST /api/bot-stop', 'api.js:287-290'],
    ['api.getDiary(filters)', 'GET /api/diary?...', 'api.js:294-300'],
    ['api.createDiaryEntry(entry)', 'POST /api/diary', 'api.js:302-309'],
    ['api.updateDiaryEntry(id, data)', 'PUT /api/diary/:id', 'api.js:311-318'],
    ['api.deleteDiaryEntry(id)', 'DELETE /api/diary/:id', 'api.js:320-323'],
    ['api.getToolboxCodes()', 'GET /api/toolbox/codes', 'api.js:328-331'],
    ['api.createToolboxCode(...)', 'POST /api/toolbox/codes', 'api.js:333-340'],
    ['api.updateToolboxCode(id, data)', 'PUT /api/toolbox/codes/:id', 'api.js:342-349'],
    ['api.deleteToolboxCode(id)', 'DELETE /api/toolbox/codes/:id', 'api.js:351-354'],
    ['api.getSefaStatus()', 'GET /api/toolbox/sefa-status', 'api.js:356-359'],
    ['api.getCnpj(cnpj, refresh)', 'GET /api/toolbox/cnpj/:cnpj', 'api.js:361-364'],
    ['api.getToolboxConfig()', 'GET /api/toolbox/config', 'api.js:366-369'],
    ['api.setToolboxConfig(key)', 'PUT /api/toolbox/config', 'api.js:371-378'],
    ['api.getNfeChave(chave)', 'GET /api/toolbox/nfe/:chave', 'api.js:380-383'],
]

add_table(
    ['Método Frontend', 'Rota Backend', 'Linha api.js'],
    api_map,
    col_widths=[5.5, 6.5, 3]
)

doc.add_heading('4.2 Socket.IO — Atualizações em tempo real', level=2)
doc.add_paragraph(
    'O Socket.IO sincroniza dados entre abas do navegador. Quando qualquer endpoint de mutação '
    '(criar, atualizar, deletar) é chamado no backend, ele emite um evento "data-changed" para todos os clientes conectados.'
)

doc.add_heading('Servidor (server/index.js:42)', level=3)
code('''const io = new Server(server, { cors: { origin: '*' } });''')

doc.add_heading('Evento de mutação (exemplo: server/index.js:115)', level=3)
code('''io.emit('data-changed');''')

doc.add_heading('Cliente (App.jsx:331-343)', level=3)
code('''useEffect(() => {
    const socket = io(window.location.origin);
    socket.on('data-changed', () => {
        loadFolders();
        loadFavorites();
        if (currentFolder) loadFiles(currentFolder);
    });
    return () => socket.disconnect();
}, [currentFolder]);''')

# ════════════════════════════════════════════════════════════
# 5. BACKEND — SERVIDOR EXPRESS
# ════════════════════════════════════════════════════════════
doc.add_heading('5. Backend — Servidor Express', level=1)

doc.add_heading('5.1 Inicialização (server/index.js:1333-1397)', level=2)
doc.add_paragraph(
    'O servidor escuta na porta 3000, bind 0.0.0.0. Exibe banner com IPs na inicialização. '
    'Suporta atalhos TTY: [R] reiniciar, [Q] sair. Encerramento graceful em SIGINT/SIGTERM.'
)

doc.add_heading('5.2 Middleware (server/index.js:63-74)', level=2)
add_table(
    ['Middleware', 'Linha', 'Função'],
    [
        ['express.json({ limit: "10mb" })', 'server/index.js:65', 'Parser de JSON com limite 10MB (para uploads base64)'],
        ['express.static("dist")', 'server/index.js:74', 'Serve os arquivos estáticos do frontend compilado'],
        ['express.static(IMAGES_DIR)', 'server/index.js:1252', 'Serve imagens de anexos em /_images'],
    ],
    col_widths=[5, 3.5, 7.5]
)

doc.add_heading('5.3 Diretório de dados (server/index.js:50-55)', level=2)
doc.add_paragraph(
    'Os dados são persistidos no sistema de arquivos. O diretório raiz é NOTION_PATH, '
    'resolvido como ../../notion/ relativo ao server/ (ou seja, C:\\Projeto\\notion\\).'
)
add_table(
    ['Dado', 'Armazenamento', 'Caminho', 'Linha'],
    [
        ['Erros (catálogo)', 'Arquivos .md', 'NOTION_PATH/<pasta>/<arquivo>.md', 'server/index.js:50'],
        ['Anexos', 'Arquivos binários', 'NOTION_PATH/_images/<pasta>/', 'server/index.js:54'],
        ['Lixeira', 'Arquivos .md movidos', 'NOTION_PATH/_erros_nao_catalogados/', 'server/index.js:55'],
        ['Favoritos', 'JSON', 'favorites.json', 'server/index.js:51'],
        ['Relatórios', 'JSON', 'reports.json', 'server/index.js:52'],
        ['Cores das pastas', 'JSON', 'folder-colors.json', 'server/index.js:53'],
        ['Diário', 'JSON', 'diary.json', 'server/index.js:30'],
        ['Códigos toolbox', 'JSON', 'toolbox.json', 'server/index.js:741'],
        ['Cache CNPJ', 'JSON', 'toolbox-cache.json', 'server/index.js:895'],
        ['Config Sintegra', 'JSON', 'toolbox-config.json', 'server/index.js:914'],
        ['Token bot', 'JSON', 'bot-config.json', 'server/index.js:29'],
    ],
    col_widths=[3.5, 3, 5, 3.5]
)

doc.add_heading('5.4 Formato dos arquivos de erro (.md)', level=2)
doc.add_paragraph('Cada arquivo .md segue um template estruturado:')
code('''# Título do Erro
Criado em: YYYY-MM-DD HH:MM

## Sistema
Nome do sistema (ex: SCGWIN)

## Contexto
Descrição do contexto em que o erro ocorre.

## Resolução (passo a passo)
1. Passo um
2. Passo dois

## Observação
Informações adicionais.

## Tags
- tag1
- tag2

## Anexos
![nome](/_images/pasta/arquivo.jpg)''')

# ════════════════════════════════════════════════════════════
# 6. ROTAS DA API
# ════════════════════════════════════════════════════════════
doc.add_heading('6. Rotas da API REST', level=1)

doc.add_paragraph(
    'O backend expõe 47 rotas REST agrupadas por recurso. '
    'Todas usam o prefixo /api (exceto /cadastrar e /api/public/*).'
)

# Pastas
doc.add_heading('6.1 Pastas', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/folders', '103', 'Lista todas as pastas (subdiretórios de NOTION_PATH)'],
        ['POST', '/api/folders', '107', 'Cria nova pasta (diretório)'],
        ['PUT', '/api/folders/:oldName', '118', 'Renomeia pasta (move todos os .md)'],
        ['DELETE', '/api/folders/:name', '135', 'Exclui pasta (move .md para lixeira)'],
    ],
    col_widths=[2, 4.5, 2, 7.5]
)

# Arquivos
doc.add_heading('6.2 Arquivos (Erros)', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/files/:folder', '173', 'Lista arquivos .md de uma pasta'],
        ['GET', '/api/file/:folder/:filename', '182', 'Lê conteúdo de um arquivo .md'],
        ['POST', '/api/file/:folder', '193', 'Cria novo arquivo .md'],
        ['PUT', '/api/file/:folder/:filename', '206', 'Atualiza conteúdo do .md'],
        ['DELETE', '/api/file/:folder/:filename', '219', 'Exclui arquivo (move para lixeira)'],
        ['PUT', '/api/file/:folder/:filename/tags', '240', 'Atualiza seção de tags no .md'],
        ['PUT', '/api/file/:folder/:filename/rename', '265', 'Renomeia arquivo .md'],
        ['PUT', '/api/file/:folder/:filename/move', '285', 'Move arquivo para outra pasta'],
    ],
    col_widths=[2, 5.5, 2, 6.5]
)

# Busca
doc.add_heading('6.3 Busca', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/search?q=...', '309', 'Busca por palavra-chave em todos os .md'],
        ['GET', '/api/search/advanced?...', '331', 'Busca avançada: texto + pasta + tags + datas'],
    ],
    col_widths=[2, 5.5, 2, 6.5]
)

# Estatísticas
doc.add_heading('6.4 Estatísticas', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/stats', '394', 'Total de erros, por pasta, recentes, por tag'],
    ],
    col_widths=[2, 4, 2, 8]
)

# Favoritos
doc.add_heading('6.5 Favoritos', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/favorites', '423', 'Lista favoritos'],
        ['POST', '/api/favorites', '427', 'Adiciona favorito'],
        ['DELETE', '/api/favorites/:folder/:filename', '438', 'Remove favorito'],
    ],
    col_widths=[2, 5.5, 2, 6.5]
)

# Tags
doc.add_heading('6.6 Tags', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/tags', '449', 'Mapa de todas as tags com arquivos associados'],
    ],
    col_widths=[2, 4, 2, 8]
)

# Lixeira
doc.add_heading('6.7 Lixeira', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/trash', '474', 'Lista arquivos na lixeira'],
        ['PUT', '/api/trash/restore/:filename', '497', 'Restaura arquivo para pasta original'],
        ['DELETE', '/api/trash/:filename', '523', 'Exclui permanentemente um arquivo'],
        ['DELETE', '/api/trash', '536', 'Esvazia toda a lixeira'],
    ],
    col_widths=[2, 5, 2, 7]
)

# Cores das pastas
doc.add_heading('6.8 Cores das Pastas', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/folder-colors', '569', 'Obtém cores de todas as pastas'],
        ['PUT', '/api/folder-colors/:folder', '574', 'Define cor de uma pasta'],
        ['DELETE', '/api/folder-colors/:folder', '585', 'Remove cor de uma pasta'],
    ],
    col_widths=[2, 5, 2, 7]
)

# Relatórios
doc.add_heading('6.9 Relatórios', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/reports', '617', 'Lista relatórios'],
        ['POST', '/api/reports', '622', 'Cria relatório'],
        ['PUT', '/api/reports/:id', '638', 'Atualiza relatório'],
        ['DELETE', '/api/reports/:id', '652', 'Exclui relatório'],
    ],
    col_widths=[2, 4, 2, 8]
)

# Diário
doc.add_heading('6.10 Diário de Ocorrências', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/diary?date=&search=', '676', 'Lista entradas (filtros opcionais: data, busca)'],
        ['POST', '/api/diary', '688', 'Cria entrada no diário'],
        ['PUT', '/api/diary/:id', '710', 'Atualiza entrada (resolvido, conteúdo, prioridade)'],
        ['DELETE', '/api/diary/:id', '729', 'Exclui entrada'],
    ],
    col_widths=[2, 5, 2, 7]
)

# Toolbox
doc.add_heading('6.11 Ferramentas (Toolbox)', level=2)

doc.add_heading('Códigos de observação', level=3)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/toolbox/codes', '754', 'Lista códigos de observação'],
        ['POST', '/api/toolbox/codes', '758', 'Cria código'],
        ['PUT', '/api/toolbox/codes/:id', '768', 'Atualiza código'],
        ['DELETE', '/api/toolbox/codes/:id', '779', 'Exclui código'],
    ],
    col_widths=[2, 4.5, 2, 7.5]
)

doc.add_heading('Status SEFAZ', level=3)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/toolbox/sefa-status', '870', 'Scraping do portal SEFAZ para disponibilidade NFe por UF (timeout 15s)'],
    ],
    col_widths=[2, 4.5, 2, 7.5]
)

doc.add_heading('Consulta CNPJ', level=3)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/toolbox/config', '927', 'Verifica se Sintegra está configurado'],
        ['PUT', '/api/toolbox/config', '931', 'Define chave de API do Sintegra'],
        ['GET', '/api/toolbox/cnpj/:cnpj', '959', 'Consulta CNPJ: ReceitaWS + BrasilAPI/Sintegra (cache 24h)'],
    ],
    col_widths=[2, 4.5, 2, 7.5]
)

doc.add_heading('Chave de acesso NFe', level=3)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/toolbox/nfe/:chave', '1076', 'Valida e decodifica chave de 44 dígitos da NFe'],
    ],
    col_widths=[2, 4.5, 2, 7.5]
)

# Anexos
doc.add_heading('6.12 Anexos', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['POST', '/api/attachments/upload', '1107', 'Upload de arquivo (base64 em JSON, limite 10MB)'],
        ['GET', '/api/attachments/:folder/:filename', '1170', 'Lista anexos referenciados no .md'],
        ['DELETE', '/api/attachments/:folder/:fileName', '1220', 'Exclui anexos e remove referências do .md'],
    ],
    col_widths=[2, 5.5, 2, 6.5]
)

# Bot
doc.add_heading('6.13 Bot Telegram', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/api/bot-config', '1258', 'Obtém status do token (mascarado)'],
        ['PUT', '/api/bot-config', '1263', 'Define/remove token do bot'],
        ['POST', '/api/bot-stop', '1326', 'Para o bot Telegram'],
    ],
    col_widths=[2, 4, 2, 8]
)

# Formulário público
doc.add_heading('6.14 Formulário Público', level=2)
add_table(
    ['Método', 'Rota', 'Linha', 'Descrição'],
    [
        ['GET', '/cadastrar', '77', 'Serve HTML standalone de cadastro'],
        ['GET', '/api/public/folders-tags', '1280', 'Lista pastas e tags disponíveis'],
        ['POST', '/api/public/submit', '1305', 'Submete novo erro (sem autenticação)'],
    ],
    col_widths=[2, 5, 2, 7]
)

# ════════════════════════════════════════════════════════════
# 7. UPLOAD DE ANEXOS
# ════════════════════════════════════════════════════════════
doc.add_heading('7. Upload de Anexos', level=1)

doc.add_paragraph(
    'Os anexos são enviados como base64 codificado no corpo da requisição JSON '
    '(não usa multipart/form-data). Isso é possível porque o body parser '
    'tem limite de 10MB (server/index.js:65).'
)

doc.add_heading('Fluxo do upload (server/index.js:1107-1167)', level=2)
steps = [
    '1. Frontend lê o arquivo como base64 (api.js:111-118)',
    '2. Envia JSON com { fileName, folder, fileContent (base64), mimeType }',
    '3. Backend decodifica o base64 em buffer binário',
    '4. Gera nome único: <timestamp>_<hex>.<extensão>',
    '5. Salva em NOTION_PATH/_images/<folder>/',
    '6. Registra metadados em _metadata.json',
    '7. Injeta referência no .md: ![nome](/_images/pasta/arquivo) ou [nome](link)',
    '8. Emite io.emit("data-changed") para atualizar clientes',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Servindo imagens (server/index.js:1252)', level=2)
code('app.use("/_images", express.static(IMAGES_DIR));')

# ════════════════════════════════════════════════════════════
# 8. BOT TELEGRAM
# ════════════════════════════════════════════════════════════
doc.add_heading('8. Bot Telegram', level=1)

doc.add_paragraph(
    'O bot Telegram (server/telegram-bot.js, 390 linhas) funciona em modo polling '
    'e permite interação completa com o catálogo de erros pelo chat.'
)

doc.add_heading('Inicialização (telegram-bot.js:121)', level=2)
code('initBot(token) — Cria instância do bot, registra comandos e handlers')

doc.add_heading('Funcionalidades', level=2)
add_table(
    ['Comando/Fluxo', 'Linha', 'Descrição'],
    [
        ['/start', 'telegram-bot.js:130', 'Registra chat para notificações, exibe menu principal'],
        ['/ajuda, /help', 'telegram-bot.js:152', 'Exibe ajuda'],
        ['/stop', 'telegram-bot.js:158', 'Para o bot'],
        ['Navegação: Pastas', 'telegram-bot.js:135-150', 'Keyboard inline para navegar pastas e ver erros'],
        ['Navegação: Buscar', 'telegram-bot.js:152', 'Busca por palavra-chave'],
        ['Navegação: Criar Erro', 'telegram-bot.js:160-200', 'Wizard multi-step: pasta → título → descrição → salva .md'],
        ['Navegação: Tags', 'telegram-bot.js:200+', 'Busca por tags via botões inline'],
    ],
    col_widths=[4, 4, 8]
)

doc.add_heading('Notificações (telegram-bot.js:367)', level=2)
doc.add_paragraph(
    'sendNotification(message) — Chamado por endpoints de mutação no server/index.js '
    'para enviar atualizações ao chat registrado.'
)

doc.add_heading('Parada (telegram-bot.js:377)', level=2)
doc.add_paragraph(
    'stopBot() — Para o polling e limpa o estado. Chamado via POST /api/bot-stop.'
)

# ════════════════════════════════════════════════════════════
# 9. TEMA
# ════════════════════════════════════════════════════════════
doc.add_heading('9. Sistema de Tema', level=1)

doc.add_paragraph(
    'O app suporta temas dark (padrão) e light. O tema é aplicado via atributo '
    'data-theme no elemento raiz (<html>) e persistido em localStorage.'
)

doc.add_heading('Variáveis CSS (global.css:1-80)', level=2)
doc.add_paragraph(
    'O tema dark define variáveis CSS em :root (ex: --bg-primary: #0f1019, --bg-secondary: #181a24). '
    'O tema light sobrescreve via [data-theme="light"] (global.css:2245+) com valores claros '
    '(ex: --bg-primary: #f8f9fc, --bg-secondary: #ffffff).'
)

doc.add_heading('Alternância (App.jsx:47)', level=2)
code('const [theme, setTheme] = useState(() => localStorage.getItem("aps-theme") || "dark");')

doc.add_heading('Aplicação (App.jsx:396)', level=2)
code('document.documentElement.setAttribute("data-theme", theme);')

# ════════════════════════════════════════════════════════════
# 10. BUILD E PRODUÇÃO
# ════════════════════════════════════════════════════════════
doc.add_heading('10. Build e Produção', level=1)

doc.add_heading('Scripts npm (package.json:6-12)', level=2)
add_table(
    ['Script', 'Comando', 'Descrição'],
    [
        ['dev', 'concurrently "npm run server" "npm run client"', 'Desenvolvimento: backend + Vite simultaneamente'],
        ['client', 'vite', 'Apenas o Vite dev server (porta 5173)'],
        ['build', 'vite build', 'Compila frontend para dist/'],
        ['server', 'node server/index.js', 'Inicia servidor de produção'],
        ['start', 'node server/index.js', 'Alias para server'],
    ],
    col_widths=[3, 6, 7]
)

doc.add_heading('Proxy de desenvolvimento (vite.config.js:8-9)', level=2)
doc.add_paragraph(
    'Em desenvolvimento, o Vite faz proxy de /api para http://localhost:3000, '
    'permitindo que frontend e backend rodem em portas diferentes.'
)
code('''server: {
    proxy: { "/api": "http://localhost:3000" }
}''')

doc.add_heading('Produção (server/index.js:74)', level=2)
doc.add_paragraph(
    'Em produção, o Express serve os arquivos estáticos de dist/ diretamente. '
    'Não há necessidade de proxy — tudo roda na porta 3000.'
)
code('app.use(express.static(path.join(__dirname, "..", "dist")));')

# ════════════════════════════════════════════════════════════
# 11. NOTAS DE ARQUITETURA
# ════════════════════════════════════════════════════════════
doc.add_heading('11. Notas de Arquitetura', level=1)

notes = [
    ('Padrão monolítico', 'App.jsx concentra todo o estado e lógica de negócio. '
     'Os componentes recebem estado e callbacks via props — não há Context API, Redux ou Zustand.'),
    ('Sem autenticação', 'Não há login, sessões, JWT ou middleware de auth. '
     'Todos os endpoints são abertos. A única proteção é mascarar o token do bot.'),
    ('Persistência em arquivos', 'Não há banco de dados. Erros são .md, metadados são JSON. '
     'Simples, mas limita escalabilidade e acesso concorrente.'),
    ('Socket.IO mínimo', 'Usado apenas para push de eventos "data-changed" — '
     'não há salas, namespaces ou autenticação no socket.'),
    ('Upload via base64', 'Anexos são enviados como base64 no JSON (não multipart). '
     'Limite de 10MB definido no body parser.'),
    ('Slash commands (server)', 'Cada rota de mutação chama io.emit("data-changed") '
     'e sendNotification() para manter clientes e Telegram sincronizados.'),
]

for title_text, desc in notes:
    bold_para(f'{title_text}: ', desc)

# ════════════════════════════════════════════════════════════
# SALVAR
# ════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), 'TECHNICAL.docx')
doc.save(output_path)
print(f'Documentação salva em: {output_path}')
